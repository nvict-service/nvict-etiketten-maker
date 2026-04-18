# -*- coding: utf-8 -*-
"""
NVict Etiketten Maker - Multi-formaat
Ontwikkeld door NVict Service

Website: www.nvict.nl
Versie: 4.0 - Multi-format Edition
"""

import sys
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image, ImageTk
import urllib.request
import urllib.error
import json
import threading
import webbrowser

# Versie informatie
CURRENT_VERSION = "4.0.0"
UPDATE_CHECK_URL = "https://nvict.nl/software/updates/etiketten_version.json"
UPDATE_CHECK_URL_FALLBACK = "https://nvict.nl/software/updates/etiketten_version.php"  # Fallback als .json geblokkeerd is

try:
    import winreg
except ImportError:
    winreg = None

def get_resource_path(relative_path):
    """Geef het absolute pad naar resource bestanden (werkt met PyInstaller)"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class Theme:
    """Bevat de kleurenschema's voor lichte en donkere thema's."""
    
    def __init__(self):
        self.is_dark_mode = self.detect_windows_theme()
        self.update_colors()
    
    def detect_windows_theme(self):
        """Detecteer Windows donker/licht thema"""
        try:
            if winreg:
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                                  r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize") as key:
                    value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
                    return value == 0
        except:
            pass
        return True
    
    def update_colors(self):
        """Update kleuren op basis van thema"""
        if self.is_dark_mode:
            self.bg_primary = "#202020"
            self.bg_secondary = "#2d2d2d"
            self.bg_tertiary = "#1a1a1a"
            self.text_primary = "#ffffff"
            self.text_secondary = "#b0b0b0"
            self.accent = "#0078d4"
            self.accent_hover = "#1084d8"
            self.border = "#3c3c3c"
            self.card_bg = "#2d2d2d"
        else:
            self.bg_primary = "#f3f3f3"
            self.bg_secondary = "#ffffff"
            self.bg_tertiary = "#e8e8e8"
            self.text_primary = "#202020"
            self.text_secondary = "#5a5a5a"
            self.accent = "#0078d4"
            self.accent_hover = "#1084d8"
            self.border = "#d1d1d1"
            self.card_bg = "#ffffff"
    
    def toggle(self):
        """Wissel tussen licht en donker thema"""
        self.is_dark_mode = not self.is_dark_mode
        self.update_colors()

def detecteer_adres_velden(kolommen):
    """Detecteer automatisch welke velden bij elkaar horen"""
    kolommen_lower = [k.lower() for k in kolommen]
    regels = []
    
    # Regel 1: Naam velden (aanhef, voorletters, achternaam)
    naam_velden = []
    naam_keywords = ['aanspr', 'aanhef', 'dhr', 'mevr', 'titel']
    voorletter_keywords = ['voorletter', 'initiaal', 'initial']
    achternaam_keywords = ['achternaam', 'naam', 'surname', 'lastname']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in naam_keywords):
            naam_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in voorletter_keywords):
            naam_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in achternaam_keywords):
            naam_velden.append(kolommen[i])
    
    if naam_velden:
        regels.append(naam_velden)
    
    # Regel 2: Extra (alleen als het bestaat)
    extra_velden = []
    extra_keywords = ['extra', 'tussenvoegsel', 'voorvoegsel', 'toevoeging']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in extra_keywords):
            extra_velden.append(kolommen[i])
    
    if extra_velden:
        regels.append(extra_velden)
    
    # Regel 3: Adres velden (straat, huisnummer)
    adres_velden = []
    straat_keywords = ['adres', 'straat', 'street', 'weg', 'laan', 'plein']
    huisnr_keywords = ['huisnummer', 'huisnr', 'nummer', 'nr', 'number']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in straat_keywords):
            adres_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in huisnr_keywords):
            adres_velden.append(kolommen[i])
    
    if adres_velden:
        regels.append(adres_velden)
    
    # Regel 4: Postcode en plaats
    postcode_velden = []
    postcode_keywords = ['postcode', 'zip', 'postal']
    plaats_keywords = ['woonplaats', 'plaats', 'stad', 'city', 'town']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in postcode_keywords):
            postcode_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in plaats_keywords):
            postcode_velden.append(kolommen[i])
    
    if postcode_velden:
        regels.append(postcode_velden)
    
    # Regel 5: Land (indien aanwezig)
    land_velden = []
    land_keywords = ['land', 'country', 'nation']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in land_keywords):
            land_velden.append(kolommen[i])
    
    if land_velden:
        regels.append(land_velden)
    
    return regels

class EtiketAanpassingWindow:
    """Window voor beperkte aanpassingen"""
    def __init__(self, parent, theme, df, excel_bestand, auto_regels):
        self.parent = parent
        self.theme = theme
        self.df = df
        self.excel_bestand = excel_bestand
        self.auto_regels = auto_regels
        self.result = None
        
        self.settings = {
            'regels': auto_regels,
            'font_grootte': 10,
            'uitlijning': 'links'
        }
        
        self.window = tk.Toplevel(parent)
        self.window.title("Etiket Preview")
        self.window.geometry("1000x800")
        self.window.minsize(900, 700)
        self.window.transient(parent)
        self.window.grab_set()
        self.window.configure(bg=theme.bg_primary)
        
        # Centreer op scherm
        self.window.update_idletasks()
        x = (self.window.winfo_screenwidth() // 2) - (1000 // 2)
        y = (self.window.winfo_screenheight() // 2) - (800 // 2)
        self.window.geometry(f"1000x800+{x}+{y}")
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup de gebruikersinterface"""
        # Header
        header_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(header_frame,
                              text="✓  Preview & Aanpassen",
                              font=("Segoe UI", 18, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(pady=20, padx=30, anchor=tk.W)
        
        # Main content
        main_container = tk.Frame(self.window, bg=self.theme.bg_primary)
        main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Links: Preview
        left_frame = tk.Frame(main_container, bg=self.theme.bg_primary)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 15))
        
        # Uitleg
        uitleg_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        uitleg_card.pack(fill=tk.X, pady=(0, 15))
        
        uitleg_content = tk.Frame(uitleg_card, bg=self.theme.card_bg)
        uitleg_content.pack(padx=25, pady=20)
        
        tk.Label(uitleg_content,
                text="📋 Automatisch gedetecteerd",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 10))
        
        # Toon welke velden op welke regel komen
        for i, regel_velden in enumerate(self.auto_regels):
            regel_text = f"Regel {i+1}: {', '.join(regel_velden)}"
            tk.Label(uitleg_content,
                    text=regel_text,
                    font=("Segoe UI", 10),
                    fg=self.theme.text_secondary,
                    bg=self.theme.card_bg,
                    anchor=tk.W).pack(fill=tk.X, pady=2)
        
        # Preview card
        preview_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        preview_card.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        preview_header = tk.Frame(preview_card, bg=self.theme.card_bg)
        preview_header.pack(fill=tk.X, padx=25, pady=(20, 15))
        
        tk.Label(preview_header,
                text="👁️ Voorbeeld",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W)
        
        tk.Label(preview_header,
                text="Eerste adres uit uw Excel (schaal 1:1)",
                font=("Segoe UI", 10),
                fg=self.theme.text_secondary,
                bg=self.theme.card_bg).pack(anchor=tk.W)
        
        # Container voor preview met vaste verhoudingen
        preview_container = tk.Frame(preview_card, bg=self.theme.card_bg)
        preview_container.pack(fill=tk.X, padx=25, pady=(10, 25))
        
        # Preview etiket met echte verhoudingen (99.1mm x 33.8mm = ongeveer 3:1)
        # Gebruiken we 360x120 pixels voor goede leesbaarheid
        preview_etiket = tk.Frame(preview_container, bg="#ffffff", 
                                 relief=tk.SOLID, borderwidth=2,
                                 width=360, height=120)
        preview_etiket.pack()
        preview_etiket.pack_propagate(False)  # Behoud vaste grootte
        
        self.preview_label = tk.Label(preview_etiket,
                                      text="",
                                      font=("Segoe UI", 9),
                                      fg="#000000",
                                      bg="#ffffff",
                                      justify=tk.LEFT,
                                      anchor=tk.NW,
                                      padx=20,
                                      pady=20)
        self.preview_label.pack(fill=tk.BOTH, expand=True)
        
        # Rechts: Opties
        right_frame = tk.Frame(main_container, bg=self.theme.bg_primary, width=280)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y)
        right_frame.pack_propagate(False)
        
        opties_card = tk.Frame(right_frame, bg=self.theme.card_bg)
        opties_card.pack(fill=tk.BOTH, expand=True)
        
        opties_content = tk.Frame(opties_card, bg=self.theme.card_bg)
        opties_content.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
        
        tk.Label(opties_content,
                text="⚙️ Opties",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 15))
        
        # Font grootte
        tk.Label(opties_content,
                text="Lettergrootte:",
                font=("Segoe UI", 11),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(5, 5))
        
        self.font_var = tk.IntVar(value=10)
        font_scale = tk.Scale(opties_content,
                             from_=8,
                             to=14,
                             orient=tk.HORIZONTAL,
                             variable=self.font_var,
                             command=lambda x: self.update_preview(),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             troughcolor=self.theme.bg_secondary,
                             highlightthickness=0)
        font_scale.pack(fill=tk.X, pady=(0, 15))
        
        # Uitlijning
        tk.Label(opties_content,
                text="Uitlijning:",
                font=("Segoe UI", 11),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(5, 5))
        
        self.align_var = tk.StringVar(value='links')
        for text, val in [("Links", "links"), ("Midden", "midden"), ("Rechts", "rechts")]:
            rb = tk.Radiobutton(opties_content,
                               text=text,
                               variable=self.align_var,
                               value=val,
                               font=("Segoe UI", 10),
                               fg=self.theme.text_primary,
                               bg=self.theme.card_bg,
                               selectcolor=self.theme.bg_secondary,
                               activebackground=self.theme.card_bg,
                               command=self.update_preview)
            rb.pack(anchor=tk.W, pady=2)
        
        # Footer met knoppen
        footer_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=85)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM)
        footer_frame.pack_propagate(False)
        
        btn_container = tk.Frame(footer_frame, bg=self.theme.bg_secondary)
        btn_container.pack(expand=True)
        
        annuleer_btn = tk.Button(btn_container,
                                text="Annuleren",
                                command=self.annuleren,
                                font=("Segoe UI", 12),
                                bg=self.theme.card_bg,
                                fg=self.theme.text_primary,
                                relief=tk.FLAT,
                                cursor="hand2",
                                padx=35,
                                pady=14)
        annuleer_btn.pack(side=tk.LEFT, padx=8)
        
        maak_btn = tk.Button(btn_container,
                           text="✓ Etiketten maken",
                           command=self.bevestigen,
                           font=("Segoe UI", 12, "bold"),
                           bg=self.theme.accent,
                           fg="#ffffff",
                           activebackground=self.theme.accent_hover,
                           activeforeground="#ffffff",
                           relief=tk.FLAT,
                           cursor="hand2",
                           padx=35,
                           pady=14)
        maak_btn.pack(side=tk.LEFT, padx=8)
        
        # Initial preview
        self.update_preview()
    
    def update_preview(self):
        """Update het preview etiket"""
        if len(self.df) == 0:
            self.preview_label.config(text="Geen data beschikbaar")
            return
        
        eerste_rij = self.df.iloc[0]
        preview_text = ""
        
        # Bouw regels
        for regel_velden in self.auto_regels:
            regel_tekst = []
            for veld in regel_velden:
                if veld in eerste_rij:
                    waarde = str(eerste_rij[veld]).strip()
                    if waarde and waarde != 'nan':
                        regel_tekst.append(waarde)
            
            if regel_tekst:
                preview_text += " ".join(regel_tekst) + "\n"
        
        # Gebruik de actuele font grootte van de slider
        font_grootte = self.font_var.get()
        self.preview_label.config(text=preview_text, font=("Segoe UI", font_grootte))
        
        # Update uitlijning
        align_map = {"links": tk.W, "midden": tk.CENTER, "rechts": tk.E}
        justify_map = {"links": tk.LEFT, "midden": tk.CENTER, "rechts": tk.RIGHT}
        self.preview_label.config(anchor=align_map[self.align_var.get()],
                                 justify=justify_map[self.align_var.get()])
    
    def annuleren(self):
        """Annuleer zonder te maken"""
        self.result = None
        self.window.destroy()
    
    def bevestigen(self):
        """Bevestig en sla instellingen op"""
        self.settings['font_grootte'] = self.font_var.get()
        self.settings['uitlijning'] = self.align_var.get()
        
        self.result = self.settings
        self.window.destroy()
    
    def wait_for_result(self):
        """Wacht tot window wordt gesloten en return result"""
        self.window.wait_window()
        return self.result

class NVictEtikettenMaker:
    # Etiket formaten configuratie
    ETIKET_FORMATEN = {
        'Herma 10825': {
            'beschrijving': '99,1 × 33,8 mm (16 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 3.38
        },
        'Herma 4625': {
            'beschrijving': '66,0 × 33,8 mm (24 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.48,
            'right_margin': 0.48,
            'etiketten_per_rij': 3,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 6.6,
            'etiket_hoogte': 3.38
        },
        'Herma 4360': {
            'beschrijving': '70,0 × 36,0 mm (24 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 0.85,
            'bottom_margin': 0.85,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 3,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 7.0,
            'etiket_hoogte': 3.6
        },
        'Herma 4267': {
            'beschrijving': '99,1 × 42,3 mm (12 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 6,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 4.23
        },
        'Herma 4425': {
            'beschrijving': '105,0 × 148,5 mm (4 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 0,
            'bottom_margin': 0,
            'left_margin': 0,
            'right_margin': 0,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 2,
            'etiket_breedte': 10.5,
            'etiket_hoogte': 14.85
        },
        'Avery L7163': {
            'beschrijving': '99,1 × 38,1 mm (14 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.3,
            'bottom_margin': 1.3,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 7,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 3.81
        }
    }
    
    def __init__(self):
        self.root = tk.Tk()
        self.theme = Theme()
        self.geselecteerd_formaat = tk.StringVar(value='Herma 10825')
        self.subtitle_label = None  # Voor dynamische update
        self.update_available = False
        self.latest_version = None
        self.download_url = None
        self.update_label = None
        self.stored_version = None  # Voor update dialog binding
        self.stored_release_notes = None  # Voor update dialog binding
        
        self.root.title("NVict Etiketten Maker")
        self.root.geometry("1000x800")
        self.root.minsize(900, 700)
        self.root.state('zoomed')  # Maximaliseer automatisch
        
        # Stel icoon in
        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except:
            pass
        
        self.setup_ui()
        self.root.configure(bg=self.theme.bg_primary)
        
    def setup_ui(self):
        """Setup de gebruikersinterface"""
        # Header
        header_frame = tk.Frame(self.root, bg=self.theme.bg_secondary, height=90)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Logo in header
        logo_container = tk.Frame(header_frame, bg=self.theme.bg_secondary)
        logo_container.pack(side=tk.LEFT, padx=25)
        
        try:
            logo_path = get_resource_path("Logo.png")
            if os.path.exists(logo_path):
                logo_img = Image.open(logo_path)
                logo_img.thumbnail((65, 65), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_img)
                logo_label = tk.Label(logo_container, image=self.logo_photo, 
                                     bg=self.theme.bg_secondary)
                logo_label.pack()
        except:
            pass
        
        # Titel in header
        title_frame = tk.Frame(header_frame, bg=self.theme.bg_secondary)
        title_frame.pack(side=tk.LEFT, padx=15, expand=True, fill=tk.BOTH)
        
        title_label = tk.Label(title_frame, 
                              text="NVict Etiketten Maker",
                              font=("Segoe UI", 22, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(anchor=tk.W, pady=(20, 0))
        
        self.subtitle_label = tk.Label(title_frame,
                                 text=self.ETIKET_FORMATEN[self.geselecteerd_formaat.get()]['beschrijving'],
                                 font=("Segoe UI", 11),
                                 fg=self.theme.text_secondary,
                                 bg=self.theme.bg_secondary)
        self.subtitle_label.pack(anchor=tk.W)
        
        # Thema toggle knop
        theme_btn = tk.Button(header_frame,
                             text="🌙" if not self.theme.is_dark_mode else "☀️",
                             command=self.toggle_theme,
                             font=("Segoe UI", 16),
                             bg=self.theme.bg_secondary,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=12)
        theme_btn.pack(side=tk.RIGHT, padx=25)
        
        # Main content
        main_frame = tk.Frame(self.root, bg=self.theme.bg_primary)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=50, pady=40)
        
        # Etiket selectie card
        select_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        select_card.pack(fill=tk.X, pady=(0, 25))
        
        select_content = tk.Frame(select_card, bg=self.theme.card_bg)
        select_content.pack(padx=35, pady=25)
        
        tk.Label(select_content,
                text="🏷️  Selecteer Etiket Formaat",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 12))
        
        # Dropdown voor etiket keuze
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('Custom.TCombobox',
                       fieldbackground=self.theme.bg_secondary,
                       background=self.theme.bg_secondary,
                       foreground=self.theme.text_primary,
                       arrowcolor=self.theme.text_primary,
                       bordercolor=self.theme.border,
                       lightcolor=self.theme.bg_secondary,
                       darkcolor=self.theme.bg_secondary,
                       borderwidth=1,
                       relief='flat')
        style.map('Custom.TCombobox',
                 fieldbackground=[('readonly', self.theme.bg_secondary)],
                 selectbackground=[('readonly', self.theme.bg_secondary)],
                 selectforeground=[('readonly', self.theme.text_primary)])
        
        combo_frame = tk.Frame(select_content, bg=self.theme.card_bg)
        combo_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.formaat_combo = ttk.Combobox(combo_frame,
                                         textvariable=self.geselecteerd_formaat,
                                         values=list(self.ETIKET_FORMATEN.keys()),
                                         state='readonly',
                                         style='Custom.TCombobox',
                                         font=("Segoe UI", 11),
                                         width=40)
        self.formaat_combo.pack(side=tk.LEFT, padx=(0, 10))
        self.formaat_combo.bind('<<ComboboxSelected>>', self.update_formaat_info)
        
        # Info over geselecteerd formaat
        self.format_info_label = tk.Label(select_content,
                                         text=f"✓ {self.ETIKET_FORMATEN[self.geselecteerd_formaat.get()]['beschrijving']}",
                                         font=("Segoe UI", 10),
                                         fg=self.theme.text_secondary,
                                         bg=self.theme.card_bg)
        self.format_info_label.pack(anchor=tk.W, pady=(5, 0))
        
        # Info card
        info_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        info_card.pack(fill=tk.X, pady=(0, 25))
        
        info_content = tk.Frame(info_card, bg=self.theme.card_bg)
        info_content.pack(padx=35, pady=30)
        
        info_icon = tk.Label(info_content,
                            text="⚡",
                            font=("Segoe UI", 40),
                            bg=self.theme.card_bg)
        info_icon.pack()
        
        info_title = tk.Label(info_content,
                             text="Excel naar Etiketten",
                             font=("Segoe UI", 18, "bold"),
                             fg=self.theme.text_primary,
                             bg=self.theme.card_bg)
        info_title.pack(pady=(12, 8))
        
        info_text = tk.Label(info_content,
                            text="Converteer uw Excel adressenlijst automatisch naar etiketten.",
                            font=("Segoe UI", 11),
                            fg=self.theme.text_secondary,
                            bg=self.theme.card_bg,
                            justify=tk.CENTER)
        info_text.pack()
        
        # Features card
        features_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        features_card.pack(fill=tk.X, pady=(0, 25))
        
        features_content = tk.Frame(features_card, bg=self.theme.card_bg)
        features_content.pack(padx=35, pady=25)
        
        tk.Label(features_content,
                text="✨ Automatisch gedetecteerd:",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 15))
        
        features = [
            ("📝 Regel 1", "Aanhef, voorletters en achternaam"),
            ("📌 Regel 2", "Extra (indien aanwezig)"),
            ("🏠 Regel 3", "Adres en huisnummer"),
            ("📍 Regel 4", "Postcode en woonplaats")
        ]
        
        for icon_text, desc in features:
            feature_frame = tk.Frame(features_content, bg=self.theme.card_bg)
            feature_frame.pack(fill=tk.X, pady=5)
            
            tk.Label(feature_frame,
                    text=icon_text,
                    font=("Segoe UI", 11, "bold"),
                    fg=self.theme.accent,
                    bg=self.theme.card_bg).pack(side=tk.LEFT, padx=(0, 10))
            
            tk.Label(feature_frame,
                    text=desc,
                    font=("Segoe UI", 11),
                    fg=self.theme.text_secondary,
                    bg=self.theme.card_bg).pack(side=tk.LEFT)
        
        # Action button
        btn_frame = tk.Frame(main_frame, bg=self.theme.bg_primary)
        btn_frame.pack(pady=25)
        
        self.start_btn = tk.Button(btn_frame,
                                   text="📁  Selecteer Excel bestand",
                                   command=self.maak_etiketten,
                                   font=("Segoe UI", 13, "bold"),
                                   bg=self.theme.accent,
                                   fg="#ffffff",
                                   activebackground=self.theme.accent_hover,
                                   activeforeground="#ffffff",
                                   relief=tk.FLAT,
                                   cursor="hand2",
                                   padx=45,
                                   pady=18)
        self.start_btn.pack()
        
        # Hover effect
        self.start_btn.bind("<Enter>", lambda e: self.start_btn.config(bg=self.theme.accent_hover))
        self.start_btn.bind("<Leave>", lambda e: self.start_btn.config(bg=self.theme.accent))
        
        # Footer
        footer_frame = tk.Frame(self.root, bg=self.theme.bg_secondary, height=50)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM)
        footer_frame.pack_propagate(False)
        
        footer_content = tk.Frame(footer_frame, bg=self.theme.bg_secondary)
        footer_content.pack(expand=True)
        
        # Versie info links
        version_label = tk.Label(footer_content,
                                text=f"v{CURRENT_VERSION}",
                                font=("Segoe UI", 9),
                                fg=self.theme.text_secondary,
                                bg=self.theme.bg_secondary)
        version_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Debug: Handmatige update check knop (alleen in debug modus)
        if os.getenv('DEBUG_MODE') == '1':
            debug_btn = tk.Label(footer_content,
                                text="🔄 Check Update",
                                font=("Segoe UI", 9, "underline"),
                                fg=self.theme.accent,
                                bg=self.theme.bg_secondary,
                                cursor="hand2")
            debug_btn.pack(side=tk.LEFT, padx=(0, 15))
            debug_btn.bind("<Button-1>", lambda e: self.manual_update_check())
        
        # Update notificatie (verborgen tot update beschikbaar is)
        self.update_label = tk.Label(footer_content,
                                     text="",
                                     font=("Segoe UI", 10, "bold", "underline"),
                                     fg=self.theme.accent,
                                     bg=self.theme.bg_secondary,
                                     cursor="hand2")
        self.update_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Main footer text
        footer_label = tk.Label(footer_content,
                               text="NVict Service  •  www.nvict.nl",
                               font=("Segoe UI", 10),
                               fg=self.theme.text_secondary,
                               bg=self.theme.bg_secondary)
        footer_label.pack(side=tk.LEFT)
        
        # Start update check in achtergrond
        threading.Thread(target=self.check_for_updates, daemon=True).start()
    
    def manual_update_check(self):
        """Handmatige update check voor debugging"""
        print("\n=== HANDMATIGE UPDATE CHECK GESTART ===")
        threading.Thread(target=self.check_for_updates, daemon=True).start()
    
    def toggle_theme(self):
        """Wissel tussen licht en donker thema"""
        self.theme.toggle()
        for widget in self.root.winfo_children():
            widget.destroy()
        self.setup_ui()
    
    def update_formaat_info(self, event=None):
        """Update de formaat info labels wanneer er een nieuw formaat wordt gekozen"""
        formaat = self.geselecteerd_formaat.get()
        specs = self.ETIKET_FORMATEN[formaat]
        
        # Update subtitle in header
        if self.subtitle_label:
            self.subtitle_label.config(text=specs['beschrijving'])
        
        # Update info label in selectie card
        if hasattr(self, 'format_info_label'):
            self.format_info_label.config(text=f"✓ {specs['beschrijving']}")
    
    def get_voorbeelden_map(self):
        """Bepaal de locatie van de voorbeelden map"""
        try:
            # Windows: Documenten\Etiketten Maker
            if os.name == 'nt':
                import winreg
                # Haal de Documenten locatie op uit registry
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                                   r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders") as key:
                    docs_folder = winreg.QueryValueEx(key, "Personal")[0]
                    voorbeelden_map = os.path.join(docs_folder, "Etiketten Maker")
            else:
                # Linux/Mac: ~/Documents/Etiketten Maker
                docs_folder = os.path.expanduser("~/Documents")
                voorbeelden_map = os.path.join(docs_folder, "Etiketten Maker")
            
            # Maak de map aan als die nog niet bestaat
            if not os.path.exists(voorbeelden_map):
                os.makedirs(voorbeelden_map, exist_ok=True)
            
            return voorbeelden_map
        except:
            # Fallback: gewoon de home directory
            return os.path.expanduser("~")
    
    def check_for_updates(self):
        """Check voor updates (draait in achtergrond thread)"""
        urls_to_try = [UPDATE_CHECK_URL, UPDATE_CHECK_URL_FALLBACK]
        
        for url_index, url in enumerate(urls_to_try):
            try:
                # Probeer versie info op te halen
                req = urllib.request.Request(
                    url,
                    headers={'User-Agent': 'NVict-Etiketten-Maker'}
                )
                
                with urllib.request.urlopen(req, timeout=5) as response:
                    data = json.loads(response.read().decode())
                    
                    latest_version = data.get('version', '')
                    download_url = data.get('download_url', '')
                    release_notes = data.get('release_notes', '')
                    
                    # Vergelijk versies
                    if self.is_newer_version(latest_version, CURRENT_VERSION):
                        print(f"[UPDATE] ✅ Update beschikbaar: v{CURRENT_VERSION} → v{latest_version}")
                        self.latest_version = latest_version
                        self.download_url = download_url
                        self.update_available = True
                        
                        # Update UI in main thread
                        self.root.after(0, self.show_update_notification, latest_version, release_notes)
                    
                    # Succesvol, stop met proberen
                    return
            
            except urllib.error.HTTPError as e:
                if e.code == 403 and url_index < len(urls_to_try) - 1:
                    continue  # Probeer fallback
                print(f"[UPDATE] ❌ HTTP Error: {e.code}")
            except urllib.error.URLError as e:
                if url_index < len(urls_to_try) - 1:
                    continue  # Probeer fallback
                print(f"[UPDATE] ❌ Kan update server niet bereiken")
            except json.JSONDecodeError:
                print(f"[UPDATE] ❌ Ongeldige server response")
            except Exception as e:
                print(f"[UPDATE] ❌ Fout: {type(e).__name__}: {e}")
    
    def is_newer_version(self, latest, current):
        """Vergelijk versie nummers"""
        try:
            latest_parts = [int(x) for x in latest.split('.')]
            current_parts = [int(x) for x in current.split('.')]
            
            # Pad met nullen als nodig
            while len(latest_parts) < 3:
                latest_parts.append(0)
            while len(current_parts) < 3:
                current_parts.append(0)
            
            return latest_parts > current_parts
        except (ValueError, AttributeError) as e:
            print(f"[UPDATE] ❌ Versie check fout: {e}")
            return False
    
    def show_update_notification(self, version, release_notes):
        """Toon update notificatie in de UI"""
        # Bewaar versie en notes als instance variabelen
        self.stored_version = version
        self.stored_release_notes = release_notes
        
        # Update het label voor als de gebruiker de dialog sluit
        if self.update_label:
            update_text = f"🔔 Update beschikbaar: v{version}"
            self.update_label.config(text=update_text)
            
            # Verwijder oude bindings eerst
            self.update_label.unbind("<Button-1>")
            
            # Bind voor als de gebruiker opnieuw wil zien
            def on_click(event):
                print(f"[UPDATE] 🖱️ Update label geklikt - toon dialog opnieuw")
                self.show_update_dialog(self.stored_version, self.stored_release_notes)
            
            self.update_label.bind("<Button-1>", on_click)
            
            print(f"[UPDATE] ✅ Update notificatie getoond: v{version}")
        
        # AUTOMATISCH de update dialog tonen
        print(f"[UPDATE] 🎯 Automatisch update dialog tonen...")
        self.show_update_dialog(version, release_notes)
    
    def show_update_dialog(self, version, release_notes):
        """Toon update dialog met details"""
        print(f"[UPDATE] 📋 Toon update dialog voor versie {version}")
        
        try:
            update_window = tk.Toplevel(self.root)
            update_window.title("Update Beschikbaar")
            update_window.geometry("600x550")
            update_window.minsize(550, 500)
            update_window.transient(self.root)
            update_window.grab_set()
            update_window.configure(bg=self.theme.bg_secondary)
            
            # Centreer
            update_window.update_idletasks()
            x = (update_window.winfo_screenwidth() // 2) - (600 // 2)
            y = (update_window.winfo_screenheight() // 2) - (550 // 2)
            update_window.geometry(f"600x550+{x}+{y}")
        except Exception as e:
            print(f"[UPDATE] ❌ Error bij aanmaken dialog: {type(e).__name__}: {e}")
            return
        
        content_frame = tk.Frame(update_window, bg=self.theme.bg_secondary)
        content_frame.pack(expand=True, fill=tk.BOTH, padx=30, pady=30)
        
        # Icon
        icon_label = tk.Label(content_frame,
                             text="🔔",
                             font=("Segoe UI", 48),
                             bg=self.theme.bg_secondary)
        icon_label.pack(pady=(0, 15))
        
        # Titel
        title_label = tk.Label(content_frame,
                              text="Update Beschikbaar!",
                              font=("Segoe UI", 16, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(pady=(0, 10))
        
        # Versie info
        version_text = f"Huidige versie: v{CURRENT_VERSION}\nNieuwe versie: v{version}"
        version_label = tk.Label(content_frame,
                                text=version_text,
                                font=("Segoe UI", 11),
                                fg=self.theme.text_secondary,
                                bg=self.theme.bg_secondary,
                                justify=tk.CENTER)
        version_label.pack(pady=(0, 15))
        
        # Release notes
        if release_notes:
            notes_frame = tk.Frame(content_frame, bg=self.theme.card_bg)
            notes_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
            
            notes_label = tk.Label(notes_frame,
                                  text="Wat is er nieuw:",
                                  font=("Segoe UI", 10, "bold"),
                                  fg=self.theme.text_primary,
                                  bg=self.theme.card_bg)
            notes_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
            
            notes_text = tk.Text(notes_frame,
                               font=("Segoe UI", 10),
                               fg=self.theme.text_secondary,
                               bg=self.theme.card_bg,
                               height=6,
                               wrap=tk.WORD,
                               relief=tk.FLAT)
            notes_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 10))
            notes_text.insert("1.0", release_notes)
            notes_text.config(state=tk.DISABLED)
        
        # Knoppen
        btn_frame = tk.Frame(content_frame, bg=self.theme.bg_secondary)
        btn_frame.pack(pady=(10, 0))
        
        download_btn = tk.Button(btn_frame,
                                text="📥 Download Update",
                                command=lambda: [webbrowser.open(self.download_url), update_window.destroy()],
                                font=("Segoe UI", 11, "bold"),
                                bg=self.theme.accent,
                                fg="#ffffff",
                                relief=tk.FLAT,
                                cursor="hand2",
                                padx=20,
                                pady=10)
        download_btn.pack(side=tk.LEFT, padx=5)
        
        later_btn = tk.Button(btn_frame,
                             text="Later",
                             command=update_window.destroy,
                             font=("Segoe UI", 11),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=20,
                             pady=10)
        later_btn.pack(side=tk.LEFT, padx=5)
    
    def maak_etiketten(self):
        """Hoofdfunctie om etiketten te maken"""
        # Bepaal start locatie voor file dialog
        start_dir = self.get_voorbeelden_map()
        
        # Selecteer Excel bestand
        excel_bestand = filedialog.askopenfilename(
            title="Selecteer Excel bestand met adressen",
            initialdir=start_dir,
            filetypes=[("Excel bestanden", "*.xlsx *.xls"), ("Alle bestanden", "*.*")]
        )
        
        if not excel_bestand:
            return
        
        try:
            # Lees Excel
            df = pd.read_excel(excel_bestand)
            
            if len(df) == 0:
                messagebox.showerror("Fout", "Het Excel bestand bevat geen gegevens.")
                return
            
            # Detecteer automatisch welke velden bij elkaar horen
            auto_regels = detecteer_adres_velden(df.columns.tolist())
            
            if not auto_regels:
                messagebox.showerror("Fout", 
                    "Kon geen adresvelden detecteren in het Excel bestand.\n\n" +
                    "Zorg ervoor dat de kolommen herkenbare namen hebben zoals:\n" +
                    "Naam, Adres, Postcode, Woonplaats, etc.")
                return
            
            # Open preview venster met beperkte opties
            preview_window = EtiketAanpassingWindow(self.root, self.theme, df, excel_bestand, auto_regels)
            settings = preview_window.wait_for_result()
            
            if settings is None:
                return
            
            # Progress window
            progress_window = tk.Toplevel(self.root)
            progress_window.title("Etiketten maken...")
            progress_window.geometry("450x180")
            progress_window.transient(self.root)
            progress_window.grab_set()
            progress_window.configure(bg=self.theme.bg_secondary)
            
            progress_window.update_idletasks()
            x = (progress_window.winfo_screenwidth() // 2) - (450 // 2)
            y = (progress_window.winfo_screenheight() // 2) - (180 // 2)
            progress_window.geometry(f"450x180+{x}+{y}")
            
            progress_label = tk.Label(progress_window,
                                     text="⏳ Etiketten worden aangemaakt...",
                                     font=("Segoe UI", 13),
                                     fg=self.theme.text_primary,
                                     bg=self.theme.bg_secondary)
            progress_label.pack(expand=True)
            progress_window.update()
            
            # Maak Word document
            doc = Document()
            
            # Haal specificaties op van geselecteerd formaat
            formaat_naam = self.geselecteerd_formaat.get()
            specs = self.ETIKET_FORMATEN[formaat_naam]
            
            # Pagina instellingen
            section = doc.sections[0]
            section.page_height = Cm(specs['page_height'])
            section.page_width = Cm(specs['page_width'])
            section.top_margin = Cm(specs['top_margin'])
            section.bottom_margin = Cm(specs['bottom_margin'])
            section.left_margin = Cm(specs['left_margin'])
            section.right_margin = Cm(specs['right_margin'])
            
            etiketten_per_rij = specs['etiketten_per_rij']
            etiketten_per_kolom = specs['etiketten_per_kolom']
            etiketten_per_vel = etiketten_per_rij * etiketten_per_kolom
            aantal_vellen = (len(df) + etiketten_per_vel - 1) // etiketten_per_vel
            
            regels_config = settings['regels']
            font_grootte = settings['font_grootte']
            uitlijning = settings['uitlijning']
            
            align_map = {
                'links': WD_ALIGN_PARAGRAPH.LEFT,
                'midden': WD_ALIGN_PARAGRAPH.CENTER,
                'rechts': WD_ALIGN_PARAGRAPH.RIGHT
            }
            word_alignment = align_map.get(uitlijning, WD_ALIGN_PARAGRAPH.LEFT)
            
            # Maak etiketten
            etiket_index = 0
            vel_nummer = 0
            
            while etiket_index < len(df):
                # Maak tabel voor dit vel
                table = doc.add_table(rows=etiketten_per_kolom, cols=etiketten_per_rij)
                table.allow_autofit = False
                
                # Verwijder spacing tussen cellen
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
                    tbl.insert(0, tblPr)
                
                # Zet cell spacing op 0
                tblCellSpacing = parse_xml(r'<w:tblCellSpacing %s w:w="0" w:type="dxa"/>' % nsdecls('w'))
                tblPr.append(tblCellSpacing)
                
                # Zet tabel layout op fixed
                tblLayout = parse_xml(r'<w:tblLayout %s w:type="fixed"/>' % nsdecls('w'))
                tblPr.append(tblLayout)
                
                # Verwijder alle borders
                tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
                tblPr.append(tblBorders)
                
                # Voor alle vellen behalve het eerste: zet page break op eerste cel
                if vel_nummer > 0:
                    first_cell = table.rows[0].cells[0]
                    first_paragraph = first_cell.paragraphs[0]
                    first_paragraph.paragraph_format.page_break_before = True
                
                vel_nummer += 1
                
                for col in table.columns:
                    col.width = Cm(specs['etiket_breedte'])
                
                for rij in range(etiketten_per_kolom):
                    table.rows[rij].height = Cm(specs['etiket_hoogte'])
                    
                    for kol in range(etiketten_per_rij):
                        if etiket_index >= len(df):
                            # Stop als we alle adressen hebben verwerkt
                            break
                        
                        row_data = df.iloc[etiket_index]
                        cell = table.rows[rij].cells[kol]
                        
                        # Zet cel margins op kleine consistente waarde (72 twips = 1.27mm)
                        tc = cell._element
                        tcPr = tc.tcPr
                        if tcPr is None:
                            tcPr = parse_xml(r'<w:tcPr %s/>' % nsdecls('w'))
                            tc.insert(0, tcPr)
                        tcMar = parse_xml(r'<w:tcMar %s><w:top w:w="72" w:type="dxa"/><w:left w:w="72" w:type="dxa"/><w:bottom w:w="36" w:type="dxa"/><w:right w:w="36" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
                        tcPr.append(tcMar)
                        
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.alignment = word_alignment
                        
                        # Bouw regels op
                        for regel_velden in regels_config:
                            regel_tekst = []
                            for veld in regel_velden:
                                if veld in row_data:
                                    waarde = str(row_data[veld]).strip()
                                    if waarde and waarde != 'nan':
                                        regel_tekst.append(waarde)
                            
                            if regel_tekst:
                                run = p.add_run(" ".join(regel_tekst) + "\n")
                                run.font.size = Pt(font_grootte)
                        
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.line_spacing = 1.0
                        cell.vertical_alignment = 1
                        
                        etiket_index += 1
                    
                    # Als we alle adressen hebben gehad, stop ook de rij-loop
                    if etiket_index >= len(df):
                        break
            
            # Opslaan
            formaat_code = formaat_naam.replace(' ', '_')
            output_bestand = excel_bestand.replace('.xlsx', f'_etiketten_{formaat_code}.docx').replace('.xls', f'_etiketten_{formaat_code}.docx')
            doc.save(output_bestand)
            
            progress_window.destroy()
            
            # Success dialog
            self.toon_success_dialog(output_bestand, len(df), aantal_vellen)
            
        except Exception as e:
            if 'progress_window' in locals():
                progress_window.destroy()
            messagebox.showerror("Fout", f"Er is een fout opgetreden:\n\n{str(e)}")
    
    def toon_success_dialog(self, output_bestand, aantal_adressen, aantal_vellen):
        """Toon success dialog"""
        success_window = tk.Toplevel(self.root)
        success_window.title("✅ Gereed")
        success_window.geometry("550x350")
        success_window.transient(self.root)
        success_window.grab_set()
        success_window.configure(bg=self.theme.bg_secondary)
        
        success_window.update_idletasks()
        x = (success_window.winfo_screenwidth() // 2) - (550 // 2)
        y = (success_window.winfo_screenheight() // 2) - (350 // 2)
        success_window.geometry(f"550x350+{x}+{y}")
        
        content_frame = tk.Frame(success_window, bg=self.theme.bg_secondary)
        content_frame.pack(expand=True, padx=45, pady=35)
        
        success_icon = tk.Label(content_frame,
                               text="✅",
                               font=("Segoe UI", 52),
                               bg=self.theme.bg_secondary)
        success_icon.pack(pady=(0, 18))
        
        success_title = tk.Label(content_frame,
                                text="Etiketten zijn aangemaakt!",
                                font=("Segoe UI", 17, "bold"),
                                fg=self.theme.text_primary,
                                bg=self.theme.bg_secondary)
        success_title.pack(pady=(0, 12))
        
        file_label = tk.Label(content_frame,
                             text=os.path.basename(output_bestand),
                             font=("Segoe UI", 11, "underline"),
                             fg=self.theme.accent,
                             bg=self.theme.bg_secondary,
                             cursor="hand2")
        file_label.pack(pady=8)
        file_label.bind("<Button-1>", lambda e: os.startfile(output_bestand))
        
        stats_frame = tk.Frame(content_frame, bg=self.theme.card_bg)
        stats_frame.pack(pady=18, fill=tk.X)
        
        stats_text = f"📊 {aantal_adressen} adressen  •  📄 {aantal_vellen} {'vel' if aantal_vellen == 1 else 'vellen'}"
        stats_label = tk.Label(stats_frame,
                              text=stats_text,
                              font=("Segoe UI", 11),
                              fg=self.theme.text_secondary,
                              bg=self.theme.card_bg,
                              padx=25,
                              pady=12)
        stats_label.pack()
        
        btn_frame = tk.Frame(content_frame, bg=self.theme.bg_secondary)
        btn_frame.pack(pady=(18, 0))
        
        open_btn = tk.Button(btn_frame,
                            text="📂 Open bestand",
                            command=lambda: [os.startfile(output_bestand), success_window.destroy()],
                            font=("Segoe UI", 11, "bold"),
                            bg=self.theme.accent,
                            fg="#ffffff",
                            relief=tk.FLAT,
                            cursor="hand2",
                            padx=25,
                            pady=10)
        open_btn.pack(side=tk.LEFT, padx=6)
        
        close_btn = tk.Button(btn_frame,
                             text="Sluiten",
                             command=success_window.destroy,
                             font=("Segoe UI", 11),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=25,
                             pady=10)
        close_btn.pack(side=tk.LEFT, padx=6)
    
    def run(self):
        """Start de applicatie"""
        self.root.mainloop()

if __name__ == "__main__":
    app = NVictEtikettenMaker()
    app.run()
