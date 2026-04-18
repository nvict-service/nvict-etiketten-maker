# -*- coding: utf-8 -*-
"""
NVict Etiketten Maker - Professional Edition
Ontwikkeld door NVict Service

Website: www.nvict.nl
Versie: 4.2.0
"""

import sys
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image, ImageTk
import urllib.request
import urllib.error
import json
import threading
import webbrowser
import tempfile
import subprocess
import platform
from datetime import datetime

# Applicatie versie
APP_VERSION = "4.2.0"
UPDATE_CHECK_URL = "https://www.nvict.nl/software/updates/etiketten_version.json"
UPDATE_CHECK_URL_FALLBACK = "https://www.nvict.nl/software/updates/etiketten_version.php"

try:
    import winreg
except ImportError:
    winreg = None

def get_resource_path(relative_path):
    """Geef het absolute pad naar resource bestanden (werkt met PyInstaller)"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class Theme:
    """Bevat de kleurenschema's voor lichte en donkere thema's."""
    
    def __init__(self):
        self.is_dark_mode = self.detect_windows_theme()
        self.update_colors()
    
    def detect_windows_theme(self):
        """Detecteer Windows donker/licht thema"""
        try:
            if winreg:
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER, 
                                  r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize") as key:
                    value, _ = winreg.QueryValueEx(key, "AppsUseLightTheme")
                    return value == 0
        except:
            pass
        return True
    
    def update_colors(self):
        """Update kleuren op basis van thema"""
        if self.is_dark_mode:
            self.bg_primary = "#202020"
            self.bg_secondary = "#2d2d2d"
            self.bg_tertiary = "#1a1a1a"
            self.text_primary = "#ffffff"
            self.text_secondary = "#b0b0b0"
            self.accent = "#0078d4"
            self.accent_hover = "#1084d8"
            self.success_color = "#107c10"
            self.border = "#3c3c3c"
            self.card_bg = "#2d2d2d"
        else:
            self.bg_primary = "#f3f3f3"
            self.bg_secondary = "#ffffff"
            self.bg_tertiary = "#e8e8e8"
            self.text_primary = "#202020"
            self.text_secondary = "#5a5a5a"
            self.accent = "#0078d4"
            self.accent_hover = "#1084d8"
            self.success_color = "#107c10"
            self.border = "#d1d1d1"
            self.card_bg = "#ffffff"
    
    def toggle(self):
        """Wissel tussen licht en donker thema"""
        self.is_dark_mode = not self.is_dark_mode
        self.update_colors()

def detecteer_adres_velden(kolommen):
    """Detecteer automatisch welke velden bij elkaar horen"""
    kolommen_lower = [k.lower() for k in kolommen]
    regels = []
    
    # Regel 1: Naam velden (aanhef, voorletters, achternaam)
    naam_velden = []
    naam_keywords = ['aanspr', 'aanhef', 'dhr', 'mevr', 'titel']
    voorletter_keywords = ['voorletter', 'initiaal', 'initial']
    achternaam_keywords = ['achternaam', 'naam', 'surname', 'lastname']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in naam_keywords):
            naam_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in voorletter_keywords):
            naam_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in achternaam_keywords):
            naam_velden.append(kolommen[i])
    
    if naam_velden:
        regels.append(naam_velden)
    
    # Regel 2: Extra (alleen als het bestaat)
    extra_velden = []
    extra_keywords = ['extra', 'tussenvoegsel', 'voorvoegsel', 'toevoeging']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in extra_keywords):
            extra_velden.append(kolommen[i])
    
    if extra_velden:
        regels.append(extra_velden)
    
    # Regel 3: Adres velden (straat, huisnummer)
    adres_velden = []
    straat_keywords = ['adres', 'straat', 'street', 'weg', 'laan', 'plein']
    huisnr_keywords = ['huisnummer', 'huisnr', 'nummer', 'nr', 'number']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in straat_keywords):
            adres_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in huisnr_keywords):
            adres_velden.append(kolommen[i])
    
    if adres_velden:
        regels.append(adres_velden)
    
    # Regel 4: Postcode en plaats
    postcode_velden = []
    postcode_keywords = ['postcode', 'zip', 'postal']
    plaats_keywords = ['woonplaats', 'plaats', 'stad', 'city', 'town']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in postcode_keywords):
            postcode_velden.append(kolommen[i])
        elif any(kw in kol_lower for kw in plaats_keywords):
            postcode_velden.append(kolommen[i])
    
    if postcode_velden:
        regels.append(postcode_velden)
    
    # Regel 5: Land (indien aanwezig)
    land_velden = []
    land_keywords = ['land', 'country', 'nation']
    
    for i, kol_lower in enumerate(kolommen_lower):
        if any(kw in kol_lower for kw in land_keywords):
            land_velden.append(kolommen[i])
    
    if land_velden:
        regels.append(land_velden)
    
    return regels

class EtiketAanpassingWindow:
    """Window voor beperkte aanpassingen"""
    def __init__(self, parent, theme, df, excel_bestand, auto_regels):
        self.parent = parent
        self.theme = theme
        self.df = df
        self.excel_bestand = excel_bestand
        self.auto_regels = auto_regels
        self.result = None
        
        self.settings = {
            'regels': auto_regels,
            'font_grootte': 10,
            'uitlijning': 'links'
        }
        
        self.window = tk.Toplevel(parent)
        self.window.title("Etiket Preview")
        self.window.geometry("1000x800")
        self.window.minsize(900, 700)
        self.window.transient(parent)
        self.window.grab_set()
        self.window.configure(bg=theme.bg_primary)
        
        # Centreer op scherm
        self.window.update_idletasks()
        x = (self.window.winfo_screenwidth() // 2) - (1000 // 2)
        y = (self.window.winfo_screenheight() // 2) - (800 // 2)
        self.window.geometry(f"1000x800+{x}+{y}")
        
        # Icoon instellen
        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                self.window.iconbitmap(icon_path)
        except:
            pass
        
        self.setup_ui()
        
    def setup_ui(self):
        """Setup de gebruikersinterface"""
        # Header
        header_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        title_label = tk.Label(header_frame,
                              text="✓  Preview & Aanpassen",
                              font=("Segoe UI", 18, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(pady=20, padx=30, anchor=tk.W)
        
        # Main content
        main_container = tk.Frame(self.window, bg=self.theme.bg_primary)
        main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        # Links: Preview
        left_frame = tk.Frame(main_container, bg=self.theme.bg_primary)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 15))
        
        # Uitleg
        uitleg_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        uitleg_card.pack(fill=tk.X, pady=(0, 15))
        
        uitleg_content = tk.Frame(uitleg_card, bg=self.theme.card_bg)
        uitleg_content.pack(padx=25, pady=20)
        
        tk.Label(uitleg_content,
                text="📋 Automatisch gedetecteerd",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 10))
        
        # Toon welke velden op welke regel komen
        for i, regel_velden in enumerate(self.auto_regels):
            regel_text = f"Regel {i+1}: {', '.join(regel_velden)}"
            tk.Label(uitleg_content,
                    text=regel_text,
                    font=("Segoe UI", 10),
                    fg=self.theme.text_secondary,
                    bg=self.theme.card_bg,
                    anchor=tk.W).pack(fill=tk.X, pady=2)
        
        # Preview card
        preview_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        preview_card.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        preview_header = tk.Frame(preview_card, bg=self.theme.card_bg)
        preview_header.pack(fill=tk.X, padx=25, pady=(20, 15))
        
        tk.Label(preview_header,
                text="👁️ Voorbeeld",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W)
        
        tk.Label(preview_header,
                text="Eerste adres uit uw Excel (schaal 1:1)",
                font=("Segoe UI", 10),
                fg=self.theme.text_secondary,
                bg=self.theme.card_bg).pack(anchor=tk.W)
        
        # Container voor preview met vaste verhoudingen
        preview_container = tk.Frame(preview_card, bg=self.theme.card_bg)
        preview_container.pack(fill=tk.X, padx=25, pady=(10, 25))
        
        # Preview etiket met echte verhoudingen (99.1mm x 33.8mm = ongeveer 3:1)
        preview_etiket = tk.Frame(preview_container, bg="#ffffff", 
                                 relief=tk.SOLID, borderwidth=2,
                                 width=360, height=120)
        preview_etiket.pack()
        preview_etiket.pack_propagate(False)
        
        self.preview_label = tk.Label(preview_etiket,
                                      text="",
                                      font=("Segoe UI", 9),
                                      fg="#000000",
                                      bg="#ffffff",
                                      justify=tk.LEFT,
                                      anchor=tk.NW,
                                      padx=20,
                                      pady=20)
        self.preview_label.pack(fill=tk.BOTH, expand=True)
        
        # Rechts: Opties
        right_frame = tk.Frame(main_container, bg=self.theme.bg_primary, width=280)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y)
        right_frame.pack_propagate(False)
        
        opties_card = tk.Frame(right_frame, bg=self.theme.card_bg)
        opties_card.pack(fill=tk.BOTH, expand=True)
        
        opties_content = tk.Frame(opties_card, bg=self.theme.card_bg)
        opties_content.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
        
        tk.Label(opties_content,
                text="⚙️ Opties",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 15))
        
        # Font grootte
        tk.Label(opties_content,
                text="Lettergrootte:",
                font=("Segoe UI", 11),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(5, 5))
        
        self.font_var = tk.IntVar(value=10)
        font_scale = tk.Scale(opties_content,
                             from_=8,
                             to=14,
                             orient=tk.HORIZONTAL,
                             variable=self.font_var,
                             command=lambda x: self.update_preview(),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             troughcolor=self.theme.bg_secondary,
                             highlightthickness=0)
        font_scale.pack(fill=tk.X, pady=(0, 15))
        
        # Uitlijning
        tk.Label(opties_content,
                text="Uitlijning:",
                font=("Segoe UI", 11),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(5, 5))
        
        self.align_var = tk.StringVar(value='links')
        for text, val in [("Links", "links"), ("Midden", "midden"), ("Rechts", "rechts")]:
            rb = tk.Radiobutton(opties_content,
                               text=text,
                               variable=self.align_var,
                               value=val,
                               font=("Segoe UI", 10),
                               fg=self.theme.text_primary,
                               bg=self.theme.card_bg,
                               selectcolor=self.theme.bg_secondary,
                               activebackground=self.theme.card_bg,
                               command=self.update_preview)
            rb.pack(anchor=tk.W, pady=2)
        
        # Footer met knoppen
        footer_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=85)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM)
        footer_frame.pack_propagate(False)
        
        btn_container = tk.Frame(footer_frame, bg=self.theme.bg_secondary)
        btn_container.pack(expand=True)
        
        annuleer_btn = tk.Button(btn_container,
                                text="Annuleren",
                                command=self.annuleren,
                                font=("Segoe UI", 12),
                                bg=self.theme.card_bg,
                                fg=self.theme.text_primary,
                                relief=tk.FLAT,
                                cursor="hand2",
                                padx=35,
                                pady=14)
        annuleer_btn.pack(side=tk.LEFT, padx=8)
        
        maak_btn = tk.Button(btn_container,
                           text="✓ Etiketten maken",
                           command=self.bevestigen,
                           font=("Segoe UI", 12, "bold"),
                           bg=self.theme.accent,
                           fg="#ffffff",
                           activebackground=self.theme.accent_hover,
                           activeforeground="#ffffff",
                           relief=tk.FLAT,
                           cursor="hand2",
                           padx=35,
                           pady=14)
        maak_btn.pack(side=tk.LEFT, padx=8)
        
        # Initial preview
        self.update_preview()
    
    def update_preview(self):
        """Update het preview etiket"""
        if len(self.df) == 0:
            self.preview_label.config(text="Geen data beschikbaar")
            return
        
        eerste_rij = self.df.iloc[0]
        preview_text = ""
        
        # Bouw regels
        for regel_velden in self.auto_regels:
            regel_tekst = []
            for veld in regel_velden:
                if veld in eerste_rij:
                    waarde = str(eerste_rij[veld]).strip()
                    if waarde and waarde != 'nan':
                        regel_tekst.append(waarde)
            
            if regel_tekst:
                preview_text += " ".join(regel_tekst) + "\n"
        
        # Gebruik de actuele font grootte van de slider
        font_grootte = self.font_var.get()
        self.preview_label.config(text=preview_text, font=("Segoe UI", font_grootte))
        
        # Update uitlijning
        align_map = {"links": tk.W, "midden": tk.CENTER, "rechts": tk.E}
        justify_map = {"links": tk.LEFT, "midden": tk.CENTER, "rechts": tk.RIGHT}
        self.preview_label.config(anchor=align_map[self.align_var.get()],
                                 justify=justify_map[self.align_var.get()])
    
    def annuleren(self):
        """Annuleer zonder te maken"""
        self.result = None
        self.window.destroy()
    
    def bevestigen(self):
        """Bevestig en sla instellingen op"""
        self.settings['font_grootte'] = self.font_var.get()
        self.settings['uitlijning'] = self.align_var.get()
        
        self.result = self.settings
        self.window.destroy()
    
    def wait_for_result(self):
        """Wacht tot window wordt gesloten en return result"""
        self.window.wait_window()
        return self.result

class EénAdresWindow:
    """Dialog voor het invoeren van één adres op alle etiketten"""
    def __init__(self, parent, theme, formaat_naam, formaat_specs):
        self.parent = parent
        self.theme = theme
        self.formaat_naam = formaat_naam
        self.formaat_specs = formaat_specs
        self.result = None

        self.etiketten_per_vel = formaat_specs['etiketten_per_rij'] * formaat_specs['etiketten_per_kolom']

        self.window = tk.Toplevel(parent)
        self.window.title("Één adres op alle etiketten")
        self.window.geometry("820x680")
        self.window.minsize(700, 600)
        self.window.transient(parent)
        self.window.grab_set()
        self.window.configure(bg=theme.bg_primary)

        self.window.update_idletasks()
        x = (self.window.winfo_screenwidth() // 2) - (820 // 2)
        y = (self.window.winfo_screenheight() // 2) - (680 // 2)
        self.window.geometry(f"820x680+{x}+{y}")

        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                self.window.iconbitmap(icon_path)
        except:
            pass

        self.setup_ui()

    def setup_ui(self):
        # Header
        header_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=70)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)

        tk.Label(header_frame,
                 text="📝  Één adres op alle etiketten",
                 font=("Segoe UI", 18, "bold"),
                 fg=self.theme.text_primary,
                 bg=self.theme.bg_secondary).pack(pady=20, padx=30, anchor=tk.W)

        # Main content
        main_container = tk.Frame(self.window, bg=self.theme.bg_primary)
        main_container.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)

        # Links: invoervelden
        left_frame = tk.Frame(main_container, bg=self.theme.bg_primary)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 15))

        # Adresregels card
        adres_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        adres_card.pack(fill=tk.X, pady=(0, 15))

        adres_content = tk.Frame(adres_card, bg=self.theme.card_bg)
        adres_content.pack(padx=20, pady=20, fill=tk.X)

        tk.Label(adres_content,
                 text="Adresregels:",
                 font=("Segoe UI", 12, "bold"),
                 fg=self.theme.text_primary,
                 bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 10))

        self.regel_vars = []
        regel_labels = [
            "Naam:",
            "Adres:",
            "Postcode + Woonplaats:",
            "Regel 4 (optioneel):",
            "Regel 5 (optioneel):"
        ]

        for label_text in regel_labels:
            row_frame = tk.Frame(adres_content, bg=self.theme.card_bg)
            row_frame.pack(fill=tk.X, pady=4)

            tk.Label(row_frame,
                     text=label_text,
                     font=("Segoe UI", 10),
                     fg=self.theme.text_secondary,
                     bg=self.theme.card_bg,
                     width=22,
                     anchor=tk.W).pack(side=tk.LEFT)

            var = tk.StringVar()
            entry = tk.Entry(row_frame,
                             textvariable=var,
                             font=("Segoe UI", 10),
                             bg=self.theme.bg_secondary,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             insertbackground=self.theme.text_primary)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=5, padx=(0, 5))
            var.trace('w', lambda *args: self.update_preview())
            self.regel_vars.append(var)

        # Aantal vellen card
        vellen_card = tk.Frame(left_frame, bg=self.theme.card_bg)
        vellen_card.pack(fill=tk.X, pady=(0, 15))

        vellen_content = tk.Frame(vellen_card, bg=self.theme.card_bg)
        vellen_content.pack(padx=20, pady=20, fill=tk.X)

        tk.Label(vellen_content,
                 text="Aantal vellen:",
                 font=("Segoe UI", 12, "bold"),
                 fg=self.theme.text_primary,
                 bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 8))

        vellen_row = tk.Frame(vellen_content, bg=self.theme.card_bg)
        vellen_row.pack(fill=tk.X)

        self.vellen_var = tk.IntVar(value=1)
        vellen_spin = tk.Spinbox(vellen_row,
                                 textvariable=self.vellen_var,
                                 from_=1, to=999,
                                 font=("Segoe UI", 12),
                                 bg=self.theme.bg_secondary,
                                 fg=self.theme.text_primary,
                                 relief=tk.FLAT,
                                 width=5,
                                 command=self.update_vellen_info)
        vellen_spin.pack(side=tk.LEFT, ipady=5)

        self.vellen_info = tk.Label(vellen_row,
                                    text=f"× {self.etiketten_per_vel} per vel = {self.etiketten_per_vel} etiketten totaal",
                                    font=("Segoe UI", 10),
                                    fg=self.theme.text_secondary,
                                    bg=self.theme.card_bg)
        self.vellen_info.pack(side=tk.LEFT, padx=12)

        # Rechts: preview + lettergrootte
        right_frame = tk.Frame(main_container, bg=self.theme.bg_primary, width=260)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y)
        right_frame.pack_propagate(False)

        preview_card = tk.Frame(right_frame, bg=self.theme.card_bg)
        preview_card.pack(fill=tk.BOTH, expand=True)

        preview_content = tk.Frame(preview_card, bg=self.theme.card_bg)
        preview_content.pack(padx=15, pady=15, fill=tk.BOTH, expand=True)

        tk.Label(preview_content,
                 text="👁️ Voorbeeld",
                 font=("Segoe UI", 12, "bold"),
                 fg=self.theme.text_primary,
                 bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 8))

        preview_etiket = tk.Frame(preview_content, bg="#ffffff",
                                  relief=tk.SOLID, borderwidth=2,
                                  width=220, height=90)
        preview_etiket.pack()
        preview_etiket.pack_propagate(False)

        self.preview_label = tk.Label(preview_etiket,
                                      text="(vul adresregels in)",
                                      font=("Segoe UI", 9),
                                      fg="#000000",
                                      bg="#ffffff",
                                      justify=tk.LEFT,
                                      anchor=tk.NW,
                                      padx=10,
                                      pady=8)
        self.preview_label.pack(fill=tk.BOTH, expand=True)

        tk.Label(preview_content,
                 text="Lettergrootte:",
                 font=("Segoe UI", 10),
                 fg=self.theme.text_primary,
                 bg=self.theme.card_bg).pack(anchor=tk.W, pady=(15, 3))

        self.font_var = tk.IntVar(value=10)
        font_scale = tk.Scale(preview_content,
                              from_=8, to=14,
                              orient=tk.HORIZONTAL,
                              variable=self.font_var,
                              command=lambda x: self.update_preview(),
                              bg=self.theme.card_bg,
                              fg=self.theme.text_primary,
                              troughcolor=self.theme.bg_secondary,
                              highlightthickness=0)
        font_scale.pack(fill=tk.X, pady=(0, 10))

        # Footer
        footer_frame = tk.Frame(self.window, bg=self.theme.bg_secondary, height=85)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM)
        footer_frame.pack_propagate(False)

        btn_container = tk.Frame(footer_frame, bg=self.theme.bg_secondary)
        btn_container.pack(expand=True)

        tk.Button(btn_container,
                  text="Annuleren",
                  command=self.annuleren,
                  font=("Segoe UI", 12),
                  bg=self.theme.card_bg,
                  fg=self.theme.text_primary,
                  relief=tk.FLAT,
                  cursor="hand2",
                  padx=35, pady=14).pack(side=tk.LEFT, padx=8)

        tk.Button(btn_container,
                  text="✓ Etiketten maken",
                  command=self.bevestigen,
                  font=("Segoe UI", 12, "bold"),
                  bg=self.theme.accent,
                  fg="#ffffff",
                  activebackground=self.theme.accent_hover,
                  activeforeground="#ffffff",
                  relief=tk.FLAT,
                  cursor="hand2",
                  padx=35, pady=14).pack(side=tk.LEFT, padx=8)

    def update_preview(self):
        lines = [v.get().strip() for v in self.regel_vars if v.get().strip()]
        preview_text = "\n".join(lines) if lines else "(vul adresregels in)"
        self.preview_label.config(text=preview_text, font=("Segoe UI", self.font_var.get()))
        self.update_vellen_info()

    def update_vellen_info(self):
        try:
            n = self.vellen_var.get()
            totaal = n * self.etiketten_per_vel
            self.vellen_info.config(
                text=f"× {self.etiketten_per_vel} per vel = {totaal} etiketten totaal"
            )
        except:
            pass

    def annuleren(self):
        self.result = None
        self.window.destroy()

    def bevestigen(self):
        regels = [v.get().strip() for v in self.regel_vars if v.get().strip()]
        if not regels:
            messagebox.showwarning("Geen adres", "Vul minimaal één adresregel in.", parent=self.window)
            return
        self.result = {
            'regels': regels,
            'aantal_vellen': self.vellen_var.get(),
            'font_grootte': self.font_var.get()
        }
        self.window.destroy()

    def wait_for_result(self):
        self.window.wait_window()
        return self.result


class NVictEtikettenMaker:
    # Etiket formaten configuratie
    ETIKET_FORMATEN = {
        'Herma 10825': {
            'beschrijving': '99,1 × 33,8 mm (16 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 3.38
        },
        'Herma 4625': {
            'beschrijving': '66,0 × 33,8 mm (24 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.48,
            'right_margin': 0.48,
            'etiketten_per_rij': 3,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 6.6,
            'etiket_hoogte': 3.38
        },
        'Herma 4360': {
            'beschrijving': '70,0 × 36,0 mm (24 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 0.85,
            'bottom_margin': 0.85,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 3,
            'etiketten_per_kolom': 8,
            'etiket_breedte': 7.0,
            'etiket_hoogte': 3.6
        },
        'Herma 4267': {
            'beschrijving': '99,1 × 42,3 mm (12 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.27,
            'bottom_margin': 1.27,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 6,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 4.23
        },
        'Herma 4425': {
            'beschrijving': '105,0 × 148,5 mm (4 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 0,
            'bottom_margin': 0,
            'left_margin': 0,
            'right_margin': 0,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 2,
            'etiket_breedte': 10.5,
            'etiket_hoogte': 14.85
        },
        'Avery L7163': {
            'beschrijving': '99,1 × 38,1 mm (14 etiketten)',
            'page_width': 21.0,
            'page_height': 29.7,
            'top_margin': 1.3,
            'bottom_margin': 1.3,
            'left_margin': 0.65,
            'right_margin': 0.65,
            'etiketten_per_rij': 2,
            'etiketten_per_kolom': 7,
            'etiket_breedte': 9.91,
            'etiket_hoogte': 3.81
        }
    }
    
    def __init__(self):
        self.root = tk.Tk()
        self.theme = Theme()
        self.geselecteerd_formaat = tk.StringVar(value='Herma 10825')
        self.subtitle_label = None
        self.update_available = False
        self.latest_version = None
        self.download_url = None
        self.update_label = None
        self.stored_version = None
        self.stored_release_notes = None
        
        self.root.title("NVict Etiketten Maker")
        self.root.geometry("1000x800")
        self.root.minsize(900, 700)
        self.root.state('zoomed')
        
        # Stel icoon in
        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
        except:
            pass
        
        self.setup_ui()
        self.root.configure(bg=self.theme.bg_primary)
        
    def setup_ui(self):
        """Setup de gebruikersinterface"""
        # Header
        header_frame = tk.Frame(self.root, bg=self.theme.bg_secondary, height=90)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        # Logo in header
        logo_container = tk.Frame(header_frame, bg=self.theme.bg_secondary)
        logo_container.pack(side=tk.LEFT, padx=25)
        
        try:
            logo_path = get_resource_path("Logo.png")
            if os.path.exists(logo_path):
                logo_img = Image.open(logo_path)
                logo_img.thumbnail((65, 65), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_img)
                logo_label = tk.Label(logo_container, image=self.logo_photo, 
                                     bg=self.theme.bg_secondary)
                logo_label.pack()
        except:
            pass
        
        # Titel in header
        title_frame = tk.Frame(header_frame, bg=self.theme.bg_secondary)
        title_frame.pack(side=tk.LEFT, padx=15, expand=True, fill=tk.BOTH)
        
        title_label = tk.Label(title_frame, 
                              text="NVict Etiketten Maker",
                              font=("Segoe UI", 22, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(anchor=tk.W, pady=(20, 0))
        
        self.subtitle_label = tk.Label(title_frame,
                                 text=self.ETIKET_FORMATEN[self.geselecteerd_formaat.get()]['beschrijving'],
                                 font=("Segoe UI", 11),
                                 fg=self.theme.text_secondary,
                                 bg=self.theme.bg_secondary)
        self.subtitle_label.pack(anchor=tk.W)
        
        # Thema toggle knop
        theme_btn = tk.Button(header_frame,
                             text="🌙" if not self.theme.is_dark_mode else "☀️",
                             command=self.toggle_theme,
                             font=("Segoe UI", 16),
                             bg=self.theme.bg_secondary,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=12)
        theme_btn.pack(side=tk.RIGHT, padx=25)
        
        # Main content
        main_frame = tk.Frame(self.root, bg=self.theme.bg_primary)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=50, pady=40)
        
        # Etiket selectie card
        select_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        select_card.pack(fill=tk.X, pady=(0, 25))
        
        select_content = tk.Frame(select_card, bg=self.theme.card_bg)
        select_content.pack(padx=35, pady=25)
        
        tk.Label(select_content,
                text="🏷️  Selecteer Etiket Formaat",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 12))
        
        # Dropdown voor etiket keuze
        style = ttk.Style()
        style.theme_use('clam')
        style.configure('Custom.TCombobox',
                       fieldbackground=self.theme.bg_secondary,
                       background=self.theme.bg_secondary,
                       foreground=self.theme.text_primary,
                       arrowcolor=self.theme.text_primary,
                       bordercolor=self.theme.border,
                       lightcolor=self.theme.bg_secondary,
                       darkcolor=self.theme.bg_secondary,
                       borderwidth=1,
                       relief='flat')
        style.map('Custom.TCombobox',
                 fieldbackground=[('readonly', self.theme.bg_secondary)],
                 selectbackground=[('readonly', self.theme.bg_secondary)],
                 selectforeground=[('readonly', self.theme.text_primary)])
        
        combo_frame = tk.Frame(select_content, bg=self.theme.card_bg)
        combo_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.formaat_combo = ttk.Combobox(combo_frame,
                                         textvariable=self.geselecteerd_formaat,
                                         values=list(self.ETIKET_FORMATEN.keys()),
                                         state='readonly',
                                         style='Custom.TCombobox',
                                         font=("Segoe UI", 11),
                                         width=40)
        self.formaat_combo.pack(side=tk.LEFT, padx=(0, 10))
        self.formaat_combo.bind('<<ComboboxSelected>>', self.update_formaat_info)
        
        # Info over geselecteerd formaat
        self.format_info_label = tk.Label(select_content,
                                         text=f"✓ {self.ETIKET_FORMATEN[self.geselecteerd_formaat.get()]['beschrijving']}",
                                         font=("Segoe UI", 10),
                                         fg=self.theme.text_secondary,
                                         bg=self.theme.card_bg)
        self.format_info_label.pack(anchor=tk.W, pady=(5, 0))
        
        # Info card
        info_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        info_card.pack(fill=tk.X, pady=(0, 25))
        
        info_content = tk.Frame(info_card, bg=self.theme.card_bg)
        info_content.pack(padx=35, pady=30)
        
        info_icon = tk.Label(info_content,
                            text="⚡",
                            font=("Segoe UI", 40),
                            bg=self.theme.card_bg)
        info_icon.pack()
        
        info_title = tk.Label(info_content,
                             text="Excel naar Etiketten",
                             font=("Segoe UI", 18, "bold"),
                             fg=self.theme.text_primary,
                             bg=self.theme.card_bg)
        info_title.pack(pady=(12, 8))
        
        info_text = tk.Label(info_content,
                            text="Converteer uw Excel adressenlijst automatisch naar etiketten.",
                            font=("Segoe UI", 11),
                            fg=self.theme.text_secondary,
                            bg=self.theme.card_bg,
                            justify=tk.CENTER)
        info_text.pack()
        
        # Features card
        features_card = tk.Frame(main_frame, bg=self.theme.card_bg)
        features_card.pack(fill=tk.X, pady=(0, 25))
        
        features_content = tk.Frame(features_card, bg=self.theme.card_bg)
        features_content.pack(padx=35, pady=25)
        
        tk.Label(features_content,
                text="✨ Automatisch gedetecteerd:",
                font=("Segoe UI", 14, "bold"),
                fg=self.theme.text_primary,
                bg=self.theme.card_bg).pack(anchor=tk.W, pady=(0, 15))
        
        features = [
            ("📝 Regel 1", "Aanhef, voorletters en achternaam"),
            ("📌 Regel 2", "Extra (indien aanwezig)"),
            ("🏠 Regel 3", "Adres en huisnummer"),
            ("📮 Regel 4", "Postcode en woonplaats")
        ]
        
        for icon_text, desc in features:
            feature_frame = tk.Frame(features_content, bg=self.theme.card_bg)
            feature_frame.pack(fill=tk.X, pady=5)
            
            tk.Label(feature_frame,
                    text=icon_text,
                    font=("Segoe UI", 11, "bold"),
                    fg=self.theme.accent,
                    bg=self.theme.card_bg).pack(side=tk.LEFT, padx=(0, 10))
            
            tk.Label(feature_frame,
                    text=desc,
                    font=("Segoe UI", 11),
                    fg=self.theme.text_secondary,
                    bg=self.theme.card_bg).pack(side=tk.LEFT)
        
        # Action buttons
        btn_frame = tk.Frame(main_frame, bg=self.theme.bg_primary)
        btn_frame.pack(pady=25)

        self.start_btn = tk.Button(btn_frame,
                                   text="📂  Selecteer Excel bestand",
                                   command=self.maak_etiketten,
                                   font=("Segoe UI", 13, "bold"),
                                   bg=self.theme.accent,
                                   fg="#ffffff",
                                   activebackground=self.theme.accent_hover,
                                   activeforeground="#ffffff",
                                   relief=tk.FLAT,
                                   cursor="hand2",
                                   padx=45,
                                   pady=18)
        self.start_btn.pack()

        self.start_btn.bind("<Enter>", lambda e: self.start_btn.config(bg=self.theme.accent_hover))
        self.start_btn.bind("<Leave>", lambda e: self.start_btn.config(bg=self.theme.accent))

        self.één_adres_btn = tk.Button(btn_frame,
                                       text="📝  Één adres op alle etiketten",
                                       command=self.maak_etiketten_één_adres,
                                       font=("Segoe UI", 11),
                                       bg=self.theme.card_bg,
                                       fg=self.theme.text_primary,
                                       activebackground=self.theme.bg_tertiary,
                                       activeforeground=self.theme.text_primary,
                                       relief=tk.FLAT,
                                       cursor="hand2",
                                       padx=45,
                                       pady=12)
        self.één_adres_btn.pack(pady=(10, 0))

        self.één_adres_btn.bind("<Enter>", lambda e: self.één_adres_btn.config(bg=self.theme.bg_tertiary))
        self.één_adres_btn.bind("<Leave>", lambda e: self.één_adres_btn.config(bg=self.theme.card_bg))
        
        # Footer
        footer_frame = tk.Frame(self.root, bg=self.theme.bg_secondary, height=50)
        footer_frame.pack(fill=tk.X, side=tk.BOTTOM)
        footer_frame.pack_propagate(False)
        
        footer_content = tk.Frame(footer_frame, bg=self.theme.bg_secondary)
        footer_content.pack(expand=True)
        
        # Versie info links
        version_label = tk.Label(footer_content,
                                text=f"v{APP_VERSION}",
                                font=("Segoe UI", 9),
                                fg=self.theme.text_secondary,
                                bg=self.theme.bg_secondary)
        version_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Update notificatie (verborgen tot update beschikbaar is)
        self.update_label = tk.Label(footer_content,
                                     text="",
                                     font=("Segoe UI", 10, "bold", "underline"),
                                     fg=self.theme.accent,
                                     bg=self.theme.bg_secondary,
                                     cursor="hand2")
        self.update_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Main footer text
        footer_label = tk.Label(footer_content,
                               text="NVict Service  •  www.nvict.nl",
                               font=("Segoe UI", 10),
                               fg=self.theme.text_secondary,
                               bg=self.theme.bg_secondary)
        footer_label.pack(side=tk.LEFT)
        
        # Start update check in achtergrond (na 2 seconden)
        self.root.after(2000, self.check_for_updates_on_startup)
    
    def check_for_updates_on_startup(self):
        """Check voor updates bij opstarten"""
        threading.Thread(target=self.check_for_updates, daemon=True).start()
    
    def check_for_updates(self):
        """Check voor updates (draait in achtergrond thread)"""
        urls_to_try = [UPDATE_CHECK_URL, UPDATE_CHECK_URL_FALLBACK]
        
        for url_index, url in enumerate(urls_to_try):
            try:
                req = urllib.request.Request(
                    url,
                    headers={'User-Agent': 'NVict-Etiketten-Maker'}
                )
                
                with urllib.request.urlopen(req, timeout=5) as response:
                    data = json.loads(response.read().decode())
                    
                    latest_version = data.get('version', '')
                    download_url = data.get('download_url', '')
                    release_notes = data.get('release_notes', '')
                    
                    # Vergelijk versies
                    if self.is_newer_version(latest_version, APP_VERSION):
                        self.latest_version = latest_version
                        self.download_url = download_url
                        self.update_available = True
                        
                        # Update UI in main thread
                        self.root.after(0, self.show_update_notification, latest_version, release_notes)
                    
                    return
            
            except urllib.error.HTTPError as e:
                if e.code == 403 and url_index < len(urls_to_try) - 1:
                    continue
            except urllib.error.URLError:
                if url_index < len(urls_to_try) - 1:
                    continue
            except:
                pass
    
    def is_newer_version(self, latest, current):
        """Vergelijk versie nummers"""
        try:
            latest_parts = [int(x) for x in latest.split('.')]
            current_parts = [int(x) for x in current.split('.')]
            
            while len(latest_parts) < 3:
                latest_parts.append(0)
            while len(current_parts) < 3:
                current_parts.append(0)
            
            return latest_parts > current_parts
        except:
            return False
    
    def show_update_notification(self, version, release_notes):
        """Toon update notificatie in de UI"""
        self.stored_version = version
        self.stored_release_notes = release_notes
        
        if self.update_label:
            update_text = f"🔔 Update beschikbaar: v{version}"
            self.update_label.config(text=update_text)
            
            self.update_label.unbind("<Button-1>")
            
            def on_click(event):
                self.show_update_dialog(self.stored_version, self.stored_release_notes)
            
            self.update_label.bind("<Button-1>", on_click)
        
        # Automatisch dialog tonen
        self.show_update_dialog(version, release_notes)
    
    def show_update_dialog(self, version, release_notes):
        """Toon update dialog met details"""
        try:
            update_window = tk.Toplevel(self.root)
            update_window.title("Update Beschikbaar")
            update_window.geometry("600x550")
            update_window.minsize(550, 500)
            update_window.transient(self.root)
            update_window.grab_set()
            update_window.configure(bg=self.theme.bg_secondary)
            
            # Icoon instellen
            try:
                icon_path = get_resource_path("favicon.ico")
                if os.path.exists(icon_path):
                    update_window.iconbitmap(icon_path)
            except:
                pass
            
            # Centreer
            update_window.update_idletasks()
            x = (update_window.winfo_screenwidth() // 2) - (600 // 2)
            y = (update_window.winfo_screenheight() // 2) - (550 // 2)
            update_window.geometry(f"600x550+{x}+{y}")
        except:
            return
        
        content_frame = tk.Frame(update_window, bg=self.theme.bg_secondary)
        content_frame.pack(expand=True, fill=tk.BOTH, padx=30, pady=30)
        
        # Icon
        icon_label = tk.Label(content_frame,
                             text="🔔",
                             font=("Segoe UI", 48),
                             bg=self.theme.bg_secondary)
        icon_label.pack(pady=(0, 15))
        
        # Titel
        title_label = tk.Label(content_frame,
                              text="Update Beschikbaar!",
                              font=("Segoe UI", 16, "bold"),
                              fg=self.theme.text_primary,
                              bg=self.theme.bg_secondary)
        title_label.pack(pady=(0, 10))
        
        # Versie info
        version_text = f"Huidige versie: v{APP_VERSION}\nNieuwe versie: v{version}"
        version_label = tk.Label(content_frame,
                                text=version_text,
                                font=("Segoe UI", 11),
                                fg=self.theme.text_secondary,
                                bg=self.theme.bg_secondary,
                                justify=tk.CENTER)
        version_label.pack(pady=(0, 15))
        
        # Release notes
        if release_notes:
            notes_frame = tk.Frame(content_frame, bg=self.theme.card_bg)
            notes_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
            
            notes_label = tk.Label(notes_frame,
                                  text="Wat is er nieuw:",
                                  font=("Segoe UI", 10, "bold"),
                                  fg=self.theme.text_primary,
                                  bg=self.theme.card_bg)
            notes_label.pack(anchor=tk.W, padx=15, pady=(10, 5))
            
            notes_text = tk.Text(notes_frame,
                               font=("Segoe UI", 10),
                               fg=self.theme.text_secondary,
                               bg=self.theme.card_bg,
                               height=6,
                               wrap=tk.WORD,
                               relief=tk.FLAT)
            notes_text.pack(fill=tk.BOTH, expand=True, padx=15, pady=(0, 10))
            notes_text.insert("1.0", release_notes)
            notes_text.config(state=tk.DISABLED)
        
        # Knoppen
        btn_frame = tk.Frame(content_frame, bg=self.theme.bg_secondary)
        btn_frame.pack(pady=(10, 0))
        
        def download_and_install():
            if self.download_url:
                update_window.destroy()
                self.download_and_install_update(self.download_url, version)
        
        def download_only():
            if self.download_url:
                webbrowser.open(self.download_url)
                update_window.destroy()
        
        download_install_btn = tk.Button(btn_frame,
                                        text="Download & Installeer",
                                        command=download_and_install,
                                        font=("Segoe UI", 11, "bold"),
                                        bg=self.theme.success_color,
                                        fg="#ffffff",
                                        relief=tk.FLAT,
                                        cursor="hand2",
                                        padx=20,
                                        pady=10)
        download_install_btn.pack(side=tk.LEFT, padx=3)
        
        download_btn = tk.Button(btn_frame,
                                text="Alleen Download",
                                command=download_only,
                                font=("Segoe UI", 11),
                                bg=self.theme.accent,
                                fg="#ffffff",
                                relief=tk.FLAT,
                                cursor="hand2",
                                padx=20,
                                pady=10)
        download_btn.pack(side=tk.LEFT, padx=3)
        
        later_btn = tk.Button(btn_frame,
                             text="Later",
                             command=update_window.destroy,
                             font=("Segoe UI", 11),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=20,
                             pady=10)
        later_btn.pack(side=tk.LEFT, padx=3)
    
    def download_and_install_update(self, download_url, version):
        """Download update en start installatie automatisch"""
        try:
            # Toon voortgang dialoog
            progress_dialog = tk.Toplevel(self.root)
            progress_dialog.title("Update Downloaden")
            progress_dialog.geometry("400x150")
            progress_dialog.configure(bg=self.theme.bg_primary)
            progress_dialog.transient(self.root)
            progress_dialog.resizable(False, False)
            
            try:
                icon_path = get_resource_path('favicon.ico')
                if os.path.exists(icon_path):
                    progress_dialog.iconbitmap(icon_path)
            except:
                pass
            
            tk.Label(progress_dialog, text="Update downloaden...", 
                    font=("Segoe UI", 12, "bold"),
                    bg=self.theme.bg_primary, 
                    fg=self.theme.text_primary).pack(pady=20)
            
            status_label = tk.Label(progress_dialog, text="Bezig met downloaden...",
                                   font=("Segoe UI", 9),
                                   bg=self.theme.bg_primary,
                                   fg=self.theme.text_secondary)
            status_label.pack(pady=10)
            
            progress_dialog.update()
            
            # Download in achtergrond thread
            def download_thread():
                try:
                    temp_dir = tempfile.gettempdir()
                    filename = f"NVict_Etiketten_v{version}_Setup.exe"
                    filepath = os.path.join(temp_dir, filename)
                    
                    urllib.request.urlretrieve(download_url, filepath)
                    
                    self.root.after(0, lambda: self._finish_download(progress_dialog, filepath))
                    
                except Exception as e:
                    self.root.after(0, lambda: self._download_error(progress_dialog, str(e)))
            
            thread = threading.Thread(target=download_thread, daemon=True)
            thread.start()
            
        except Exception as e:
            messagebox.showerror("Download Fout", 
                f"Kan update niet downloaden:\n{str(e)}\n\nProbeer handmatig te downloaden via de website.")
    
    def _finish_download(self, progress_dialog, filepath):
        """Voltooi download en start installer"""
        try:
            progress_dialog.destroy()
            
            if os.path.exists(filepath):
                if messagebox.askyesno("Update Download Voltooid",
                    f"Update succesvol gedownload!\n\n"
                    f"Wilt u de installer nu starten?\n\n"
                    f"Let op: Sluit eerst NVict Etiketten Maker af voordat u de installatie voltooit."):
                    
                    if platform.system() == "Windows":
                        os.startfile(filepath)
                    elif platform.system() == "Darwin":
                        subprocess.run(["open", filepath])
                    else:
                        subprocess.run(["xdg-open", filepath])
                    
                    self.root.after(1000, self.root.destroy)
            else:
                messagebox.showerror("Fout", "Download bestand niet gevonden")
                
        except Exception as e:
            messagebox.showerror("Fout", f"Kan installer niet starten:\n{str(e)}")
    
    def _download_error(self, progress_dialog, error_msg):
        """Toon download fout"""
        progress_dialog.destroy()
        messagebox.showerror("Download Fout",
            f"Kan update niet downloaden:\n{error_msg}\n\n"
            f"Probeer handmatig te downloaden via de website.")
    
    def toggle_theme(self):
        """Wissel tussen licht en donker thema"""
        self.theme.toggle()
        for widget in self.root.winfo_children():
            widget.destroy()
        self.setup_ui()
    
    def update_formaat_info(self, event=None):
        """Update de formaat info labels wanneer er een nieuw formaat wordt gekozen"""
        formaat = self.geselecteerd_formaat.get()
        specs = self.ETIKET_FORMATEN[formaat]
        
        if self.subtitle_label:
            self.subtitle_label.config(text=specs['beschrijving'])
        
        if hasattr(self, 'format_info_label'):
            self.format_info_label.config(text=f"✓ {specs['beschrijving']}")
    
    def get_voorbeelden_map(self):
        """Bepaal de locatie van de voorbeelden map"""
        try:
            if os.name == 'nt':
                import winreg
                with winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                                   r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders") as key:
                    docs_folder = winreg.QueryValueEx(key, "Personal")[0]
                    voorbeelden_map = os.path.join(docs_folder, "Etiketten Maker")
            else:
                docs_folder = os.path.expanduser("~/Documents")
                voorbeelden_map = os.path.join(docs_folder, "Etiketten Maker")
            
            if not os.path.exists(voorbeelden_map):
                os.makedirs(voorbeelden_map, exist_ok=True)
            
            return voorbeelden_map
        except:
            return os.path.expanduser("~")
    
    def maak_etiketten(self):
        """Hoofdfunctie om etiketten te maken"""
        start_dir = self.get_voorbeelden_map()
        
        excel_bestand = filedialog.askopenfilename(
            title="Selecteer Excel bestand met adressen",
            initialdir=start_dir,
            filetypes=[("Excel bestanden", "*.xlsx *.xls"), ("Alle bestanden", "*.*")]
        )
        
        if not excel_bestand:
            return

        _ok = None
        try:
            df = pd.read_excel(excel_bestand)
            
            if len(df) == 0:
                messagebox.showerror("Fout", "Het Excel bestand bevat geen gegevens.")
                return
            
            auto_regels = detecteer_adres_velden(df.columns.tolist())
            
            if not auto_regels:
                messagebox.showerror("Fout", 
                    "Kon geen adresvelden detecteren in het Excel bestand.\n\n" +
                    "Zorg ervoor dat de kolommen herkenbare namen hebben zoals:\n" +
                    "Naam, Adres, Postcode, Woonplaats, etc.")
                return
            
            preview_window = EtiketAanpassingWindow(self.root, self.theme, df, excel_bestand, auto_regels)
            settings = preview_window.wait_for_result()
            
            if settings is None:
                return
            
            # Progress window
            progress_window = tk.Toplevel(self.root)
            progress_window.title("Etiketten maken...")
            progress_window.geometry("450x180")
            progress_window.transient(self.root)
            progress_window.grab_set()
            progress_window.configure(bg=self.theme.bg_secondary)
            
            try:
                icon_path = get_resource_path("favicon.ico")
                if os.path.exists(icon_path):
                    progress_window.iconbitmap(icon_path)
            except:
                pass
            
            progress_window.update_idletasks()
            x = (progress_window.winfo_screenwidth() // 2) - (450 // 2)
            y = (progress_window.winfo_screenheight() // 2) - (180 // 2)
            progress_window.geometry(f"450x180+{x}+{y}")
            
            progress_label = tk.Label(progress_window,
                                     text="⏳ Etiketten worden aangemaakt...",
                                     font=("Segoe UI", 13),
                                     fg=self.theme.text_primary,
                                     bg=self.theme.bg_secondary)
            progress_label.pack(expand=True)
            progress_window.update()
            
            # Maak Word document
            doc = Document()
            
            formaat_naam = self.geselecteerd_formaat.get()
            specs = self.ETIKET_FORMATEN[formaat_naam]
            
            # Pagina instellingen
            section = doc.sections[0]
            section.page_height = Cm(specs['page_height'])
            section.page_width = Cm(specs['page_width'])
            section.top_margin = Cm(specs['top_margin'])
            section.bottom_margin = Cm(specs['bottom_margin'])
            section.left_margin = Cm(specs['left_margin'])
            section.right_margin = Cm(specs['right_margin'])
            
            etiketten_per_rij = specs['etiketten_per_rij']
            etiketten_per_kolom = specs['etiketten_per_kolom']
            etiketten_per_vel = etiketten_per_rij * etiketten_per_kolom
            aantal_vellen = (len(df) + etiketten_per_vel - 1) // etiketten_per_vel
            
            regels_config = settings['regels']
            font_grootte = settings['font_grootte']
            uitlijning = settings['uitlijning']
            
            align_map = {
                'links': WD_ALIGN_PARAGRAPH.LEFT,
                'midden': WD_ALIGN_PARAGRAPH.CENTER,
                'rechts': WD_ALIGN_PARAGRAPH.RIGHT
            }
            word_alignment = align_map.get(uitlijning, WD_ALIGN_PARAGRAPH.LEFT)
            
            # Maak etiketten
            etiket_index = 0
            vel_nummer = 0
            
            while etiket_index < len(df):
                table = doc.add_table(rows=etiketten_per_kolom, cols=etiketten_per_rij)
                table.allow_autofit = False
                
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
                    tbl.insert(0, tblPr)
                
                tblCellSpacing = parse_xml(r'<w:tblCellSpacing %s w:w="0" w:type="dxa"/>' % nsdecls('w'))
                tblPr.append(tblCellSpacing)
                
                tblLayout = parse_xml(r'<w:tblLayout %s w:type="fixed"/>' % nsdecls('w'))
                tblPr.append(tblLayout)
                
                tblBorders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
                tblPr.append(tblBorders)
                
                if vel_nummer > 0:
                    first_cell = table.rows[0].cells[0]
                    first_paragraph = first_cell.paragraphs[0]
                    first_paragraph.paragraph_format.page_break_before = True
                
                vel_nummer += 1
                
                for col in table.columns:
                    col.width = Cm(specs['etiket_breedte'])
                
                for rij in range(etiketten_per_kolom):
                    table.rows[rij].height = Cm(specs['etiket_hoogte'])
                    
                    for kol in range(etiketten_per_rij):
                        if etiket_index >= len(df):
                            break
                        
                        row_data = df.iloc[etiket_index]
                        cell = table.rows[rij].cells[kol]
                        
                        tc = cell._element
                        tcPr = tc.tcPr
                        if tcPr is None:
                            tcPr = parse_xml(r'<w:tcPr %s/>' % nsdecls('w'))
                            tc.insert(0, tcPr)
                        tcMar = parse_xml(r'<w:tcMar %s><w:top w:w="72" w:type="dxa"/><w:left w:w="72" w:type="dxa"/><w:bottom w:w="36" w:type="dxa"/><w:right w:w="36" w:type="dxa"/></w:tcMar>' % nsdecls('w'))
                        tcPr.append(tcMar)
                        
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.alignment = word_alignment
                        
                        for regel_velden in regels_config:
                            regel_tekst = []
                            for veld in regel_velden:
                                if veld in row_data:
                                    waarde = str(row_data[veld]).strip()
                                    if waarde and waarde != 'nan':
                                        regel_tekst.append(waarde)
                            
                            if regel_tekst:
                                run = p.add_run(" ".join(regel_tekst) + "\n")
                                run.font.size = Pt(font_grootte)
                        
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.line_spacing = 1.0
                        cell.vertical_alignment = 1
                        
                        etiket_index += 1
                    
                    if etiket_index >= len(df):
                        break
            
            # Opslaan
            formaat_code = formaat_naam.replace(' ', '_')
            output_bestand = excel_bestand.replace('.xlsx', f'_etiketten_{formaat_code}.docx').replace('.xls', f'_etiketten_{formaat_code}.docx')
            doc.save(output_bestand)
            progress_window.destroy()
            _ok = (output_bestand, len(df), aantal_vellen)

        except Exception as e:
            if 'progress_window' in locals():
                try:
                    progress_window.destroy()
                except:
                    pass
            _ok = None
            messagebox.showerror("Fout", f"Er is een fout opgetreden:\n\n{str(e)}")

        if _ok:
            self.toon_success_dialog(*_ok)
    
    def maak_etiketten_één_adres(self):
        """Maak etiketten met één adres herhaald op alle posities"""
        formaat_naam = self.geselecteerd_formaat.get()
        specs = self.ETIKET_FORMATEN[formaat_naam]

        dialog = EénAdresWindow(self.root, self.theme, formaat_naam, specs)
        result = dialog.wait_for_result()

        if result is None:
            return

        regels = result['regels']
        aantal_vellen = result['aantal_vellen']
        font_grootte = result['font_grootte']

        etiketten_per_rij = specs['etiketten_per_rij']
        etiketten_per_kolom = specs['etiketten_per_kolom']
        etiketten_per_vel = etiketten_per_rij * etiketten_per_kolom
        totaal_etiketten = aantal_vellen * etiketten_per_vel

        # Vraag opslaglocatie
        start_dir = self.get_voorbeelden_map()
        formaat_code = formaat_naam.replace(' ', '_')

        output_bestand = filedialog.asksaveasfilename(
            title="Sla etiketten op als...",
            initialdir=start_dir,
            initialfile=f"etiketten_{formaat_code}.docx",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("Alle bestanden", "*.*")]
        )

        if not output_bestand:
            return

        # Progress window
        progress_window = tk.Toplevel(self.root)
        progress_window.title("Etiketten maken...")
        progress_window.geometry("450x180")
        progress_window.transient(self.root)
        progress_window.grab_set()
        progress_window.configure(bg=self.theme.bg_secondary)

        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                progress_window.iconbitmap(icon_path)
        except:
            pass

        progress_window.update_idletasks()
        x = (progress_window.winfo_screenwidth() // 2) - (450 // 2)
        y = (progress_window.winfo_screenheight() // 2) - (180 // 2)
        progress_window.geometry(f"450x180+{x}+{y}")

        tk.Label(progress_window,
                 text="⏳ Etiketten worden aangemaakt...",
                 font=("Segoe UI", 13),
                 fg=self.theme.text_primary,
                 bg=self.theme.bg_secondary).pack(expand=True)
        progress_window.update()

        _ok = None
        try:
            doc = Document()

            section = doc.sections[0]
            section.page_height = Cm(specs['page_height'])
            section.page_width = Cm(specs['page_width'])
            section.top_margin = Cm(specs['top_margin'])
            section.bottom_margin = Cm(specs['bottom_margin'])
            section.left_margin = Cm(specs['left_margin'])
            section.right_margin = Cm(specs['right_margin'])

            etiket_index = 0
            vel_nummer = 0

            while etiket_index < totaal_etiketten:
                table = doc.add_table(rows=etiketten_per_kolom, cols=etiketten_per_rij)
                table.allow_autofit = False

                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
                    tbl.insert(0, tblPr)

                tblPr.append(parse_xml(r'<w:tblCellSpacing %s w:w="0" w:type="dxa"/>' % nsdecls('w')))
                tblPr.append(parse_xml(r'<w:tblLayout %s w:type="fixed"/>' % nsdecls('w')))
                tblPr.append(parse_xml(
                    r'<w:tblBorders %s>'
                    r'<w:top w:val="none"/><w:left w:val="none"/>'
                    r'<w:bottom w:val="none"/><w:right w:val="none"/>'
                    r'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
                    r'</w:tblBorders>' % nsdecls('w')
                ))

                if vel_nummer > 0:
                    table.rows[0].cells[0].paragraphs[0].paragraph_format.page_break_before = True

                vel_nummer += 1

                for col in table.columns:
                    col.width = Cm(specs['etiket_breedte'])

                for rij in range(etiketten_per_kolom):
                    table.rows[rij].height = Cm(specs['etiket_hoogte'])

                    for kol in range(etiketten_per_rij):
                        if etiket_index >= totaal_etiketten:
                            break

                        cell = table.rows[rij].cells[kol]

                        tc = cell._element
                        tcPr = tc.tcPr
                        if tcPr is None:
                            tcPr = parse_xml(r'<w:tcPr %s/>' % nsdecls('w'))
                            tc.insert(0, tcPr)
                        tcPr.append(parse_xml(
                            r'<w:tcMar %s>'
                            r'<w:top w:w="72" w:type="dxa"/><w:left w:w="72" w:type="dxa"/>'
                            r'<w:bottom w:w="36" w:type="dxa"/><w:right w:w="36" w:type="dxa"/>'
                            r'</w:tcMar>' % nsdecls('w')
                        ))

                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        for regel in regels:
                            run = p.add_run(regel + "\n")
                            run.font.size = Pt(font_grootte)

                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.line_spacing = 1.0
                        cell.vertical_alignment = 1

                        etiket_index += 1

                    if etiket_index >= totaal_etiketten:
                        break

            doc.save(output_bestand)
            progress_window.destroy()
            _ok = (output_bestand, totaal_etiketten, aantal_vellen)

        except Exception as e:
            try:
                progress_window.destroy()
            except:
                pass
            _ok = None
            messagebox.showerror("Fout", f"Er is een fout opgetreden:\n\n{str(e)}")

        if _ok:
            self.toon_success_dialog(*_ok)

    def toon_success_dialog(self, output_bestand, aantal_adressen, aantal_vellen):
        """Toon success dialog"""
        success_window = tk.Toplevel(self.root)
        success_window.title("✅ Gereed")
        success_window.geometry("550x350")
        success_window.transient(self.root)
        success_window.grab_set()
        success_window.configure(bg=self.theme.bg_secondary)
        
        try:
            icon_path = get_resource_path("favicon.ico")
            if os.path.exists(icon_path):
                success_window.iconbitmap(icon_path)
        except:
            pass
        
        success_window.update_idletasks()
        x = (success_window.winfo_screenwidth() // 2) - (550 // 2)
        y = (success_window.winfo_screenheight() // 2) - (350 // 2)
        success_window.geometry(f"550x350+{x}+{y}")
        
        content_frame = tk.Frame(success_window, bg=self.theme.bg_secondary)
        content_frame.pack(expand=True, padx=45, pady=35)
        
        success_icon = tk.Label(content_frame,
                               text="✅",
                               font=("Segoe UI", 52),
                               bg=self.theme.bg_secondary)
        success_icon.pack(pady=(0, 18))
        
        success_title = tk.Label(content_frame,
                                text="Etiketten zijn aangemaakt!",
                                font=("Segoe UI", 17, "bold"),
                                fg=self.theme.text_primary,
                                bg=self.theme.bg_secondary)
        success_title.pack(pady=(0, 12))
        
        file_label = tk.Label(content_frame,
                             text=os.path.basename(output_bestand),
                             font=("Segoe UI", 11, "underline"),
                             fg=self.theme.accent,
                             bg=self.theme.bg_secondary,
                             cursor="hand2")
        file_label.pack(pady=8)
        file_label.bind("<Button-1>", lambda e: os.startfile(output_bestand))
        
        stats_frame = tk.Frame(content_frame, bg=self.theme.card_bg)
        stats_frame.pack(pady=18, fill=tk.X)
        
        stats_text = f"📊 {aantal_adressen} adressen  •  📄 {aantal_vellen} {'vel' if aantal_vellen == 1 else 'vellen'}"
        stats_label = tk.Label(stats_frame,
                              text=stats_text,
                              font=("Segoe UI", 11),
                              fg=self.theme.text_secondary,
                              bg=self.theme.card_bg,
                              padx=25,
                              pady=12)
        stats_label.pack()
        
        btn_frame = tk.Frame(content_frame, bg=self.theme.bg_secondary)
        btn_frame.pack(pady=(18, 0))
        
        open_btn = tk.Button(btn_frame,
                            text="📂 Open bestand",
                            command=lambda: [os.startfile(output_bestand), success_window.destroy()],
                            font=("Segoe UI", 11, "bold"),
                            bg=self.theme.accent,
                            fg="#ffffff",
                            relief=tk.FLAT,
                            cursor="hand2",
                            padx=25,
                            pady=10)
        open_btn.pack(side=tk.LEFT, padx=6)
        
        close_btn = tk.Button(btn_frame,
                             text="Sluiten",
                             command=success_window.destroy,
                             font=("Segoe UI", 11),
                             bg=self.theme.card_bg,
                             fg=self.theme.text_primary,
                             relief=tk.FLAT,
                             cursor="hand2",
                             padx=25,
                             pady=10)
        close_btn.pack(side=tk.LEFT, padx=6)
    
    def run(self):
        """Start de applicatie"""
        self.root.mainloop()

if __name__ == "__main__":
    app = NVictEtikettenMaker()
    app.run()
